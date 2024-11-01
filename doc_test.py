import docx
from docx.shared import Inches
from docx2pdf import convert

# pip install python-docx
# pip install docx2pdf

json_cv = {
  "Email": "talentconnect.be@madello.com",
  "Highest_Qualification_Held": "System & Network Engineering and Business Skills",
  "Skill_Set": [
    "Network infrastructure",
    "Firewall architecture",
    "Cybersecurity best practices",
    "VMware NSX-T",
    "Linux/Windows Server administration",
    "Okta",
    "Duo",
    "Palo Alto Networks",
    "Fortinet",
    "Cisco",
    "VPN configurations",
    "SSL decryption",
    "SASE solutions"
  ],
  "Profile_Summary": "A highly motivated and self-taught Network & Security Engineer with proven expertise in network infrastructure, firewall architecture, and cybersecurity best practices. Brings hands-on experience from top-tier organizations, contributing to complex migration projects, hybrid cloud integrations, and high-availability firewall deployments. Demonstrates strong capabilities in designing and managing secure, scalable infrastructures, while also authoring technical content on modern network security topics.",
  "Current_Employer": "Angel One Orange Cyberdefense",
  "Street": "",
  "Zip_Code": "",
  "Experience_in_Years": 4,
  "LinkedIn__s": "",
  "City": "",
  "State": "",
  "Country": "Belgium",
  "Secondary_Email": "",
  "Current_Job_Title": "Network Security Engineer / SOC Engineer",
  "Salutation": "",
  "First_Name": "Ömer",
  "Full_Name": "Ömer Kurnaz",
  "Last_Name": "Kurnaz",
  "Mobile": "+32 479 26 61 58",
  "Current_Salary": "",
  "Experience_Details": [
    {
      "Company": "Angel One Orange Cyberdefense",
      "I_currently_work_here": True,
      "Summary": "Monitor and respond to security-related incidents as part of a SOC team. Provide 2nd-line support and handle escalation cases. Consult clients on secure design, architecture, methodology, and best practices.",
      "Work_Duration": {
        "from": "March 2024",
        "to": "Present"
      },
      "Occupation_Title": "Network Security Engineer / SOC Engineer"
    },
    {
      "Company": "NTT",
      "I_currently_work_here": False,
      "Summary": "Led migration projects from FortiGate to Palo Alto Networks Firewall & Panorama for clients including Erasme and RSVZ (Inasti). Delivered Prisma Access SASE and SD-WAN implementations.",
      "Work_Duration": {
        "from": "November 2022",
        "to": "October 2023"
      },
      "Occupation_Title": "Network Security Engineer"
    },
    {
      "Company": "Eurosys",
      "I_currently_work_here": False,
      "Summary": "Provided infrastructure support across firewalls, switches, Wi-Fi, and virtualization. Worked with VMware vSphere, Veeam backup solutions, and Windows Server.",
      "Work_Duration": {
        "from": "March 2022",
        "to": "October 2022"
      },
      "Occupation_Title": "Network Security Engineer"
    },
    {
      "Company": "Faatech.be",
      "I_currently_work_here": True,
      "Summary": "Created technical documentation and in-depth guides on firewall technologies, VPNs, SSL decryption, load balancers, identity management, and virtualization.",
      "Work_Duration": {
        "from": "April 2020",
        "to": "Present"
      },
      "Occupation_Title": "Technical Writer & Researcher"
    }
  ],
  "Phone": "",
  "Lang_Proficiency": [
    "Dutch Native",
    "English Professional Proficiency"
  ],
  "Educational_Details": [
    {
      "Institute_School": "SyntraPXL, Genk",
      "Currently_pursuing": False,
      "Degree": "System & Network Engineering and Business Skills",
      "Major_Department": "",
      "Duration": {
        "from": "2019",
        "to": "2020"
      }
    }
  ]
}


print (json_cv["Email"])
# Create a document
doc = docx.Document()
doc.add_picture('madello.png', width=Inches(1.25))


# Add a paragraph to the document
p = doc.add_paragraph()

# Add some formatting to the paragraph
p.paragraph_format.line_spacing = 1
p.paragraph_format.space_after = 0

# Add a run to the paragraph
run = p.add_run("python-docx")

# Add some formatting to the run
run.bold = True
run.italic = True
run.font.name = 'Arial'
run.font.size = docx.shared.Pt(16)

# Add more text to the same paragraph
run = p.add_run(" Tutorial")

# Format the run
run.bold = True
run.font.name = 'Arial'
run.font.size = docx.shared.Pt(16)

# Add another paragraph (left blank for an empty line)
doc.add_paragraph()

# Add another paragraph
p = doc.add_paragraph()

# Add a run and format it
run = p.add_run("This is my first python-docx tutorial!")
run.font.name = 'Arial'
run.font.size = docx.shared.Pt(12)

# Save the document
doc.save("resume.docx")


import subprocess
import os

# Function to convert PPT/PPTX to PDF on Linux
libreoffice_path = '/opt/libreoffice7.1/program/soffice'
subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', "/home/muralikmani/resume.docx", '--outdir', os.path.dirname("/home/muralikmani/resume.pdf")])



